from datetime import datetime, timedelta
from typing import Optional
import os
import secrets

from docx import Document
import PyPDF2
from fastapi import (
    FastAPI,
    File,
    Form,
    UploadFile,
    HTTPException,
    Depends,
    status,
    Request
)
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session

# Local imports
import schemas
import models
from database import get_db
from models import (
    Application,
    JobDescription,
    Resume,
    Session as SessionModel,
    User,
    User as UserModel,
    Base
)
from schemas import (
    JobDescriptionCreate,
    ResumeCreate,
    Token,
    UserCreate,
    UserProfile,
    User as UserSchema
)
from crud import (
    create_job_description,
    create_resume,
    get_all_resumes,
    get_job_description
)
from utils import (
    ACCESS_TOKEN_EXPIRE_MINUTES,
    compare_skills,
    create_access_token,
    extract_job_description_info,
    extract_resume_info,
    get_current_active_user,
    get_current_user,
    get_password_hash,
    match_resume_to_job,
    verify_password
)  # Make sure these are imported

# Initialize FastAPI app
app = FastAPI()

# Ensure uploads directory exists
os.makedirs("uploads", exist_ok=True)

# Mount static files for uploads
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)




# File upload endpoint
@app.post("/upload-resume/")
async def upload_resume(
    file: UploadFile = File(...), 
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Save the uploaded file
    file_location = f"uploads/{file.filename}"
    with open(file_location, "wb") as buffer:
        buffer.write(await file.read())

    # Read the file content
    if file.filename.endswith(".pdf"):
        with open(file_location, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            content = "".join(page.extract_text() for page in reader.pages)
    elif file.filename.endswith(".docx"):
        import docx
        doc = docx.Document(file_location)
        content = "\n".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        try:
            with open(file_location, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_location, "r", encoding="latin-1") as f:
                content = f.read()
    else:
        return {"error": "Unsupported file format"}

    # Extract resume info
    skills, experience, education = extract_resume_info(content)
    
    # Convert skills list to comma-separated string for storage
    skills_str = ", ".join(skills) if isinstance(skills, list) else skills

    # Save to database with user_id
    db_resume = Resume(
        filename=file.filename, 
        content=content,
        skills=skills_str,
        experience=experience,
        education=education,
        user_id=current_user.id
    )
    db.add(db_resume)
    db.commit()
    db.refresh(db_resume)

    return {
        "id": db_resume.id,
        "filename": file.filename, 
        "skills": skills, 
        "experience": experience, 
        "education": education
    }

@app.post("/upload-job-description/")
async def upload_job_description(
    file: UploadFile = File(...), 
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Save the uploaded file
    file_location = f"uploads/{file.filename}"
    with open(file_location, "wb") as buffer:
        buffer.write(await file.read())

    # Read the file content
    if file.filename.endswith(".pdf"):
        with open(file_location, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            content = "".join(page.extract_text() for page in reader.pages)
    elif file.filename.endswith(".docx"):
        import docx
        doc = docx.Document(file_location)
        content = "\n".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        try:
            with open(file_location, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_location, "r", encoding="latin-1") as f:
                content = f.read()
    else:
        return {"error": "Unsupported file format"}

    # Save to database with user_id
    db_job_description = JobDescription(
        filename=file.filename,
        content=content,
        user_id=current_user.id
    )
    db.add(db_job_description)
    db.commit()
    db.refresh(db_job_description)

    return {
        "id": db_job_description.id,
        "filename": file.filename, 
        "title": file.filename.replace(".pdf", "").replace(".docx", "").replace(".txt", "").replace("_", " "),
        "upload_date": db_job_description.upload_date
    }

@app.get("/rank-candidates/")
async def rank_candidates(
    job_description_id: int, 
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user)
):
    # Retrieve the job description
    job_description = db.query(JobDescription).filter(JobDescription.id == job_description_id).first()
    if not job_description:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    # Verify the job belongs to this recruiter
    if job_description.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="You can only view candidates for your own job postings")

    job_skills = extract_job_description_info(job_description.content)

    # Only get resumes from candidates who applied to THIS job
    applications = db.query(Application).filter(Application.job_id == job_description_id).all()
    
    if not applications:
        return {"ranked_candidates": [], "message": "No applications yet for this job"}

    ranked_candidates = []
    for app in applications:
        resume = app.resume
        if not resume:
            continue
            
        # Extract skills from the resume
        resume_skills, experience, education = extract_resume_info(resume.content)

        # Compare skills to find matches
        matching_skills = compare_skills(resume_skills, job_skills)

        # Calculate match score (e.g., based on the number of matching skills)
        match_score = len(matching_skills) / len(job_skills) if job_skills else 0

        ranked_candidates.append({
            "resume_id": resume.id,
            "filename": resume.filename,
            "applicant_email": app.user.email if app.user else "Unknown",
            "application_id": app.id,
            "application_status": app.status,
            "match_score": match_score,
            "matching_skills": matching_skills,
        })

    # Sort candidates by match score (descending order)
    ranked_candidates.sort(key=lambda x: x["match_score"], reverse=True)

    return {"ranked_candidates": ranked_candidates}

# New endpoint to list all job descriptions (for current recruiter)
@app.get("/job-descriptions/")
def get_job_descriptions(
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user)
):
    # Only return jobs belonging to the current recruiter
    jobs = db.query(JobDescription).filter(JobDescription.user_id == current_user.id).all()
    return [{
        "id": job.id,
        "filename": job.filename,
        "title": job.filename.replace(".pdf", "").replace(".docx", "").replace(".txt", "").replace("_", " "),
        "upload_date": job.upload_date
    } for job in jobs]


# Delete job description endpoint
@app.delete("/job-descriptions/{job_id}")
def delete_job_description(
    job_id: int,
    db: Session = Depends(get_db),
    current_user: UserModel = Depends(get_current_user)
):
    # Check if user is a recruiter
    if current_user.role != "recruiter":
        raise HTTPException(status_code=403, detail="Only recruiters can delete job postings")
    
    # Find the job description
    job = db.query(JobDescription).filter(JobDescription.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job description not found")
    
    # Check if the job belongs to this recruiter (if user_id is set)
    if job.user_id and job.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="You can only delete your own job postings")
    
    # Delete the file from uploads folder
    file_path = os.path.join("uploads", job.filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Error deleting file: {e}")
    
    # Delete from database
    db.delete(job)
    db.commit()
    
    return {"message": "Job description deleted successfully"}


# Auth endpoints
@app.post("/register", response_model=UserSchema)
def register_user(user_data: UserCreate, db: Session = Depends(get_db)):
    hashed_pw = get_password_hash(user_data.password)
    user = UserModel(
        email=user_data.email,
        role=user_data.role,
        hashed_password=hashed_pw,
        is_active=True,
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user

@app.post("/login", response_model=Token)  # Change path to "/login"
async def login_for_access_token(
    request: Request,
    form_data: OAuth2PasswordRequestForm = Depends(), 
    db: Session = Depends(get_db)
):
    user = db.query(UserModel).filter(UserModel.email == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=401,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.email, "role": user.role}, 
        expires_delta=access_token_expires
    )
    
    # Create session record
    session = SessionModel(
        user_id=user.id,
        token=access_token,
        expires_at=datetime.utcnow() + access_token_expires,
        device_info=request.headers.get("user-agent", "Unknown"),
        ip_address=request.client.host if request.client else None
    )
    db.add(session)
    db.commit()
    
    # Return both token and role
    return {
        "access_token": access_token, 
        "token_type": "bearer",
        "role": user.role  # Explicitly include role
    }


@app.post("/logout")
async def logout(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
    token: str = Depends(models.oauth2_scheme) if hasattr(models, 'oauth2_scheme') else None
):
    """Invalidate the current session"""
    from utils import oauth2_scheme
    # Get token from header and invalidate it
    # We need to get the token differently since it's already validated
    sessions = db.query(SessionModel).filter(
        SessionModel.user_id == current_user.id,
        SessionModel.is_active == True
    ).all()
    
    for session in sessions:
        session.is_active = False
    
    db.commit()
    return {"message": "Successfully logged out"}


@app.post("/logout-current")
async def logout_current_session(
    request: Request,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Invalidate only the current session (current device)"""
    auth_header = request.headers.get("authorization", "")
    if auth_header.startswith("Bearer "):
        token = auth_header[7:]
        session = db.query(SessionModel).filter(
            SessionModel.token == token,
            SessionModel.is_active == True
        ).first()
        if session:
            session.is_active = False
            db.commit()
    return {"message": "Successfully logged out from current device"}


@app.get("/sessions")
async def get_active_sessions(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all active sessions for the current user"""
    sessions = db.query(SessionModel).filter(
        SessionModel.user_id == current_user.id,
        SessionModel.is_active == True,
        SessionModel.expires_at > datetime.utcnow()
    ).all()
    
    return [
        {
            "id": session.id,
            "created_at": session.created_at,
            "expires_at": session.expires_at,
            "device_info": session.device_info,
            "ip_address": session.ip_address
        }
        for session in sessions
    ]


@app.delete("/sessions/{session_id}")
async def revoke_session(
    session_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Revoke a specific session (logout from a specific device)"""
    session = db.query(SessionModel).filter(
        SessionModel.id == session_id,
        SessionModel.user_id == current_user.id
    ).first()
    
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session.is_active = False
    db.commit()
    return {"message": "Session revoked successfully"}


@app.get("/users/me", response_model=schemas.UserProfile)
async def read_users_me(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)  # Include this if needed for additional queries
):
    return current_user

@app.put("/users/me", response_model=schemas.UserProfile)
async def update_user_me(
    profile_picture: Optional[UploadFile] = File(None),
    current_password: Optional[str] = Form(None),
    new_password: Optional[str] = Form(None),
    current_user: models.User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    # Handle profile picture update
    if profile_picture:
        # Create uploads directory if it doesn't exist
        os.makedirs("uploads/profile_pictures", exist_ok=True)
        
        # Generate unique filename
        file_location = f"uploads/profile_pictures/user_{current_user.id}_{profile_picture.filename}"
        
        # Save file
        with open(file_location, "wb") as buffer:
            buffer.write(await profile_picture.read())
        
        # Update profile picture path
        current_user.profile_picture = f"/{file_location}"
    
    # Handle password change
    if current_password and new_password:
        # Verify current password
        if not verify_password(current_password, current_user.hashed_password):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Current password is incorrect"
            )
        
        # Update to new password
        current_user.hashed_password = get_password_hash(new_password)
    
    # Commit changes to database
    db.commit()
    db.refresh(current_user)
    
    return current_user

@app.get("/resumes/me")
async def get_user_resumes(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    resumes = db.query(Resume).filter(Resume.user_id == current_user.id).all()
    return [{
        "id": resume.id,
        "filename": resume.filename,
        "skills": resume.skills,
        "experience": resume.experience,
        "education": resume.education
    } for resume in resumes]


@app.delete("/resumes/{resume_id}")
async def delete_resume(
    resume_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Find the resume
    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Check if the resume belongs to the current user
    if resume.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="You can only delete your own resumes")
    
    # Delete the file from uploads folder
    file_path = os.path.join("uploads", resume.filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Error deleting file: {e}")
    
    # Delete from database
    db.delete(resume)
    db.commit()
    
    return {"message": "Resume deleted successfully"}


@app.get("/applications")
async def get_user_applications(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all applications for the current user"""
    applications = db.query(Application).filter(
        Application.user_id == current_user.id
    ).all()
    
    return [{
        "id": app.id,
        "job_id": app.job_id,
        "job_title": app.job.filename.replace(".pdf", "").replace(".docx", "").replace(".txt", "").replace("_", " ") if app.job else "Unknown",
        "resume_filename": app.resume.filename if app.resume else "Unknown",
        "status": app.status,
        "applied_date": app.applied_date
    } for app in applications]


@app.post("/applications")
async def create_application(
    job_id: int = Form(...),
    resume_id: int = Form(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Apply to a job with a resume"""
    # Check if job exists
    job = db.query(JobDescription).filter(JobDescription.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Check if resume exists and belongs to user
    resume = db.query(Resume).filter(Resume.id == resume_id, Resume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found or doesn't belong to you")
    
    # Check if already applied
    existing = db.query(Application).filter(
        Application.user_id == current_user.id,
        Application.job_id == job_id
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail="You have already applied to this job")
    
    # Create application
    application = Application(
        user_id=current_user.id,
        job_id=job_id,
        resume_id=resume_id,
        status="pending"
    )
    db.add(application)
    db.commit()
    db.refresh(application)
    
    return {
        "id": application.id,
        "job_id": application.job_id,
        "job_title": job.filename.replace(".pdf", "").replace(".docx", "").replace(".txt", "").replace("_", " "),
        "resume_filename": resume.filename,
        "status": application.status,
        "applied_date": application.applied_date,
        "message": "Application submitted successfully!"
    }


@app.get("/jobs/available")
async def get_available_jobs(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all available job postings for applicants to browse"""
    jobs = db.query(JobDescription).all()
    return [{
        "id": job.id,
        "title": job.filename.replace(".pdf", "").replace(".docx", "").replace(".txt", "").replace("_", " "),
        "filename": job.filename,
        "upload_date": job.upload_date,
        "skills_required": extract_job_description_info(job.content) if job.content else []
    } for job in jobs]


@app.get("/stats/recruiter")
async def get_recruiter_stats(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get stats for recruiter dashboard"""
    if current_user.role != "recruiter":
        raise HTTPException(status_code=403, detail="Access denied. Recruiters only.")
    
    # Count unique candidates who applied to this recruiter's jobs
    total_candidates = db.query(Application.user_id).join(JobDescription).filter(
        JobDescription.user_id == current_user.id
    ).distinct().count()
    
    total_jobs = db.query(JobDescription).filter(JobDescription.user_id == current_user.id).count()
    total_applications = db.query(Application).join(JobDescription).filter(
        JobDescription.user_id == current_user.id
    ).count()
    
    return {
        "total_candidates": total_candidates,
        "active_jobs": total_jobs,
        "total_applications": total_applications
    }


@app.get("/recruiter/applications")
async def get_recruiter_applications(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all applications for jobs posted by this recruiter"""
    if current_user.role != "recruiter":
        raise HTTPException(status_code=403, detail="Access denied. Recruiters only.")
    
    applications = db.query(Application).join(JobDescription).filter(
        JobDescription.user_id == current_user.id
    ).all()
    
    return [{
        "id": app.id,
        "job_id": app.job_id,
        "job_title": app.job.filename.replace(".pdf", "").replace(".docx", "").replace(".txt", "").replace("_", " ") if app.job else "Unknown",
        "applicant_email": app.user.email if app.user else "Unknown",
        "applicant_id": app.user_id,
        "resume_id": app.resume_id,
        "resume_filename": app.resume.filename if app.resume else "Unknown",
        "resume_skills": app.resume.skills if app.resume else "",
        "status": app.status,
        "applied_date": app.applied_date
    } for app in applications]


@app.get("/recruiter/applications/{job_id}")
async def get_job_applications(
    job_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all applications for a specific job"""
    if current_user.role != "recruiter":
        raise HTTPException(status_code=403, detail="Access denied. Recruiters only.")
    
    # Verify job belongs to this recruiter
    job = db.query(JobDescription).filter(
        JobDescription.id == job_id,
        JobDescription.user_id == current_user.id
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found or doesn't belong to you")
    
    applications = db.query(Application).filter(Application.job_id == job_id).all()
    
    return [{
        "id": app.id,
        "applicant_email": app.user.email if app.user else "Unknown",
        "applicant_id": app.user_id,
        "resume_id": app.resume_id,
        "resume_filename": app.resume.filename if app.resume else "Unknown",
        "resume_skills": app.resume.skills if app.resume else "",
        "status": app.status,
        "applied_date": app.applied_date
    } for app in applications]


@app.put("/recruiter/applications/{application_id}")
async def update_application_status(
    application_id: int,
    status: str = Form(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Accept or decline an application"""
    if current_user.role != "recruiter":
        raise HTTPException(status_code=403, detail="Access denied. Recruiters only.")
    
    if status not in ["pending", "reviewed", "accepted", "rejected"]:
        raise HTTPException(status_code=400, detail="Invalid status. Use: pending, reviewed, accepted, or rejected")
    
    # Get the application and verify it belongs to this recruiter's job
    application = db.query(Application).join(JobDescription).filter(
        Application.id == application_id,
        JobDescription.user_id == current_user.id
    ).first()
    
    if not application:
        raise HTTPException(status_code=404, detail="Application not found or doesn't belong to your job posting")
    
    application.status = status
    db.commit()
    db.refresh(application)
    
    return {
        "id": application.id,
        "status": application.status,
        "message": f"Application {status} successfully"
    }