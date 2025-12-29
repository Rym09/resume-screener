from datetime import timedelta
from typing import Optional
import os

from docx import Document
import PyPDF2
from fastapi import (
    FastAPI,
    File,
    Form,
    UploadFile,
    HTTPException,
    Depends,
    status
)
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session

# Local imports
import schemas
import models
from database import get_db
from models import (
    JobDescription,
    Resume,
    User,
    User as UserModel,
    Base
)
from schemas import (
    JobDescriptionCreate,
    ResumeCreate,
    Token,
    UserCreate,
    UserProfile,
    User as UserSchema
)
from crud import (
    create_job_description,
    create_resume,
    get_all_resumes,
    get_job_description
)
from utils import (
    ACCESS_TOKEN_EXPIRE_MINUTES,
    compare_skills,
    create_access_token,
    extract_job_description_info,
    extract_resume_info,
    get_current_active_user,
    get_current_user,
    get_password_hash,
    match_resume_to_job,
    verify_password
)  # Make sure these are imported# Initialize FastAPI app
app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)




# File upload endpoint
@app.post("/upload-resume/")
async def upload_resume(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Save the uploaded file
    file_location = f"uploads/{file.filename}"
    with open(file_location, "wb") as buffer:
        buffer.write(await file.read())

    # Read the file content
    if file.filename.endswith(".pdf"):
        with open(file_location, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            content = "".join(page.extract_text() for page in reader.pages)
    elif file.filename.endswith(".docx"):
        import docx
        doc = docx.Document(file_location)
        content = "\n".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        try:
            with open(file_location, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_location, "r", encoding="latin-1") as f:
                content = f.read()
    else:
        return {"error": "Unsupported file format"}

    # Extract resume info
    skills, experience, education = extract_resume_info(content)

    # Save to database
    resume_data = ResumeCreate(filename=file.filename, content=content)
    db_resume = create_resume(db, resume_data)
    db_resume.skills = skills
    db_resume.experience = experience
    db_resume.education = education
    db.commit()

    return {"filename": file.filename, "skills": skills, "experience": experience, "education": education}

@app.post("/upload-job-description/")
async def upload_job_description(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Save the uploaded file
    file_location = f"uploads/{file.filename}"
    with open(file_location, "wb") as buffer:
        buffer.write(await file.read())

    # Read the file content
    if file.filename.endswith(".pdf"):
        with open(file_location, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            content = "".join(page.extract_text() for page in reader.pages)
    elif file.filename.endswith(".docx"):
        import docx
        doc = docx.Document(file_location)
        content = "\n".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        try:
            with open(file_location, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_location, "r", encoding="latin-1") as f:
                content = f.read()
    else:
        return {"error": "Unsupported file format"}

    # Save to database
    job_description_data = JobDescriptionCreate(filename=file.filename, content=content)
    db_job_description = JobDescription(**job_description_data.dict())
    db.add(db_job_description)
    db.commit()
    db.refresh(db_job_description)

    return {"filename": file.filename, "content": content}

@app.get("/rank-candidates/")
async def rank_candidates(job_description_id: int, db: Session = Depends(get_db)):
    # Retrieve the job description
    job_description = db.query(JobDescription).filter(JobDescription.id == job_description_id).first()
    if not job_description:
        raise HTTPException(status_code=404, detail="Job description not found")

    job_skills = extract_job_description_info(job_description.content)

    resumes = db.query(Resume).all()

    ranked_candidates = []
    for resume in resumes:
        # Extract skills from the resume
        resume_skills, experience, education = extract_resume_info(resume.content)

        # Compare skills to find matches
        matching_skills = compare_skills(resume_skills, job_skills)

        # Calculate match score (e.g., based on the number of matching skills)
        match_score = len(matching_skills) / len(job_skills) if job_skills else 0

        ranked_candidates.append({
            "resume_id": resume.id,
            "filename": resume.filename,
            "match_score": match_score,
            "matching_skills": matching_skills,
        })

    # Sort candidates by match score (descending order)
    ranked_candidates.sort(key=lambda x: x["match_score"], reverse=True)

    return {"ranked_candidates": ranked_candidates}

# New endpoint to list all job descriptions
@app.get("/job-descriptions/")
def get_job_descriptions(db: Session = Depends(get_db)):
    jobs = db.query(JobDescription).all()
    return [{
        "id": job.id,
        "title": job.filename.replace(".pdf", "").replace(".docx", "").replace("_", " "),
        "upload_date": job.upload_date  # Add this field to your model
    } for job in jobs]


# Auth endpoints
@app.post("/register", response_model=UserSchema)
def register_user(user_data: UserCreate, db: Session = Depends(get_db)):
    hashed_pw = get_password_hash(user_data.password)
    user = UserModel(
        email=user_data.email,
        role=user_data.role,
        hashed_password=hashed_pw,
        is_active=True,
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user

@app.post("/login", response_model=Token)  # Change path to "/login"
async def login_for_access_token(
    form_data: OAuth2PasswordRequestForm = Depends(), 
    db: Session = Depends(get_db)
):
    user = db.query(UserModel).filter(UserModel.email == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=401,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.email, "role": user.role}, 
        expires_delta=access_token_expires
    )
    
    # Return both token and role
    return {
        "access_token": access_token, 
        "token_type": "bearer",
        "role": user.role  # Explicitly include role
    }


@app.get("/users/me", response_model=schemas.UserProfile)
async def read_users_me(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)  # Include this if needed for additional queries
):
    return current_user
# In your main.py (backend)
@app.put("/users/me", response_model=schemas.UserProfile)
async def update_user_me(
    profile_picture: UploadFile = File(None),
    current_user: models.User = Depends(get_current_active_user),  # Fixed this line
    db: Session = Depends(get_db)
):
    if profile_picture:
        # Create uploads directory if it doesn't exist
        os.makedirs("uploads/profile_pictures", exist_ok=True)
        
        # Generate unique filename
        file_location = f"uploads/profile_pictures/user_{current_user.id}_{profile_picture.filename}"
        
        # Save file
        with open(file_location, "wb") as buffer:
            buffer.write(await profile_picture.read())
        
        # Update user in database
        current_user.profile_picture = f"/{file_location}"
        db.commit()
        db.refresh(current_user)
    
    return current_user