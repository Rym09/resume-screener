from docx import Document
from fastapi import FastAPI, File, HTTPException, UploadFile, Depends
from sqlalchemy.orm import Session
from models import JobDescription, Resume, SessionLocal, Base
from schemas import JobDescriptionCreate, ResumeCreate  # Ensure this import is present
from crud import create_job_description, create_resume, get_all_resumes, get_job_description
from utils import compare_skills, extract_job_description_info, extract_resume_info, match_resume_to_job
import os
import PyPDF2
from fastapi.middleware.cors import CORSMiddleware

# Initialize FastAPI app
app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Allow your React app's origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# Dependency to get the database session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# File upload endpoint
@app.post("/upload-resume/")
async def upload_resume(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Save the uploaded file
    file_location = f"uploads/{file.filename}"
    with open(file_location, "wb") as buffer:
        buffer.write(await file.read())

    # Read the file content
    if file.filename.endswith(".pdf"):
        with open(file_location, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            content = "".join(page.extract_text() for page in reader.pages)
    elif file.filename.endswith(".docx"):
        import docx
        doc = docx.Document(file_location)
        content = "\n".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        try:
            with open(file_location, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_location, "r", encoding="latin-1") as f:
                content = f.read()
    else:
        return {"error": "Unsupported file format"}

    # Extract resume info
    skills, experience, education = extract_resume_info(content)

    # Save to database
    resume_data = ResumeCreate(filename=file.filename, content=content)
    db_resume = create_resume(db, resume_data)
    db_resume.skills = skills
    db_resume.experience = experience
    db_resume.education = education
    db.commit()

    return {"filename": file.filename, "skills": skills, "experience": experience, "education": education}

@app.post("/upload-job-description/")
async def upload_job_description(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Save the uploaded file
    file_location = f"uploads/{file.filename}"
    with open(file_location, "wb") as buffer:
        buffer.write(await file.read())

    # Read the file content
    if file.filename.endswith(".pdf"):
        with open(file_location, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            content = "".join(page.extract_text() for page in reader.pages)
    elif file.filename.endswith(".docx"):
        import docx
        doc = docx.Document(file_location)
        content = "\n".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        try:
            with open(file_location, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_location, "r", encoding="latin-1") as f:
                content = f.read()
    else:
        return {"error": "Unsupported file format"}

    # Save to database
    job_description_data = JobDescriptionCreate(filename=file.filename, content=content)
    db_job_description = JobDescription(**job_description_data.dict())
    db.add(db_job_description)
    db.commit()
    db.refresh(db_job_description)

    return {"filename": file.filename, "content": content}

@app.get("/rank-candidates/")
async def rank_candidates(job_description_id: int, db: Session = Depends(get_db)):
    # Retrieve the job description
    job_description = db.query(JobDescription).filter(JobDescription.id == job_description_id).first()
    if not job_description:
        raise HTTPException(status_code=404, detail="Job description not found")

    # Extract skills from the job description
    job_skills = extract_job_description_info(job_description.content)

    # Retrieve all resumes
    resumes = db.query(Resume).all()

    # Calculate match scores and matching skills
    ranked_candidates = []
    for resume in resumes:
        # Extract skills from the resume
        resume_skills, experience, education = extract_resume_info(resume.content)

        # Compare skills to find matches
        matching_skills = compare_skills(resume_skills, job_skills)

        # Calculate match score (e.g., based on the number of matching skills)
        match_score = len(matching_skills) / len(job_skills) if job_skills else 0

        ranked_candidates.append({
            "resume_id": resume.id,
            "filename": resume.filename,
            "match_score": match_score,
            "matching_skills": matching_skills,
        })

    # Sort candidates by match score (descending order)
    ranked_candidates.sort(key=lambda x: x["match_score"], reverse=True)

    return {"ranked_candidates": ranked_candidates}


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()